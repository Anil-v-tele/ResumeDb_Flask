import docx
from docx.api import Document as Docs
import os
import win32com
from win32com import client
from PyPDF2 import PdfReader
from io import StringIO
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
import pandas as pd
import pytesseract as pt
import pdf2image
import os
import json

##======================================================
#### Function to read word file with doc extension 
## return filename, file text
##======================================================

def docxReader(filePath,link):
    file_name = file_text = ""
    try:
        if not filePath.split(".")[1] == "docx":
            print(f'-- File format passed is not correct this function only take .doc extensions')
            return        
        file_name = os.path.basename(filePath)
        doc = docx.Document(filePath)
        document = Docs(filePath)
        table = document.tables[0]
        
        data = []
        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            # if i == 0:
            #     keys = tuple(text)
            #     continue

            # Construct a dictionary for this row, mapping
            # keys to values for this row
            # row_data = dict(zip(keys, text))
            # data.append(row_data)
            for each_word in text:
                data.append(each_word)
        
        for para in doc.paragraphs:
            data.append(para.text)

        # paras = [para.append(p.text) for p in doc.paragraphs if p.text]    
        "print(f'=== Output type is a {type(paras)} of {type(paras[1])} \ntotal length is {len(paras)} ===')"
        print(f'-- File {file_name} , readed successfully having length of {len(paras)} ')
        file_text = "/n".join(data)
    except Exception as e:
        print(e)
    if file_name and filePath and file_text:
        return {"FileName":file_name,"FilePath":filePath,"FileText":file_text, "Link":link}
    else:
        return

# data_dct = docxReader("C:/Projects/Import_PDF_Word_Python/Sample_File_DOCX.docx")
# print(json.dumps(data_dct,indent=4))


def docReader(filePath,link):
    file_name = file_text = ""
    try:
        if not filePath.split(".")[1] == "doc":
            print(f'-- File format passed is not correct this function only take .doc extensions')
            return        
        file_name = os.path.basename(filePath)    
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False        
        _ = word.Documents.Open(filePath)        
        doc = word.ActiveDocument
        file_text = doc.Range().text    
        doc.Close()
        word.Quit()
    except Exception as e:
        print(e)    
    if file_name and filePath and file_text:
        return {"FileName":file_name,"FilePath":filePath,"FileText":file_text,"Link":link}
    else:
        return

# doctext = docReader("C:/Users/vedan/Downloads/Salman_Khalid.doc")
# print(doctext)

def pdfReader(filePath,link):
    file_name = file_text = ""
    try:
        if not filePath.split(".")[1] == "pdf":
            print(f'-- File format passed is not correct this function only take .pdf extensions')
            return        
        file_name = os.path.basename(filePath)
        with open(filePath, 'rb') as f:
            pdf = PdfReader(f)
            number_of_pages = len(pdf.pages)
            text_lst = []
            for i in range(number_of_pages):
                page = pdf.pages[i]
                text = page.extract_text()
                text_lst.append(text)
                file_text = "\n".join(text_lst)                   
        f.close()        
    except Exception as e:
        print(e)        
    if file_name and filePath and file_text:
        return {"FileName":file_name,"FilePath":filePath,"FileText":file_text,"Link":link
        }
    else:
        return


def parse_resume(filePath,link):
    file_extension=  filePath.split(".")[1]
    pyld = {}
    match file_extension:
        case "doc":
            pyld = docReader(filePath,link)
        case "docx":
            pyld = docxReader(filePath,link)
        case "pdf":
            pyld = pdfReader(filePath,link)    
    return pyld




# data_dct = pdfReader("C:/Users/vedan/Downloads/Resume-Anil-Kumar.pdf")
# print(json.dumps(data_dct,indent=4))